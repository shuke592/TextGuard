"""
TextGuard 文档处理服务
支持 .doc / .docx / .pdf / .txt 文件的文本提取与校对结果回写
"""
import os
import re
import subprocess
import shutil
import tempfile
from html import escape as html_escape
from typing import List, Tuple
from loguru import logger

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.shared import RGBColor


def extract_text_from_file(file_path: str, file_ext: str) -> str:
    """
    从文件中提取纯文本

    :param file_path: 文件路径
    :param file_ext: 文件扩展名（含点，如 .docx）
    :return: 提取的文本内容
    """
    ext = file_ext.lower()
    if ext == ".txt":
        return _extract_txt(file_path)
    elif ext == ".doc":
        return _extract_doc(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".pdf":
        return _extract_pdf(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _extract_txt(file_path: str) -> str:
    """提取 TXT 文件文本"""
    encodings = ["utf-8", "gbk", "gb2312", "utf-16", "latin-1"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError("无法识别文件编码，请确保文件为 UTF-8 或 GBK 编码")


def _extract_doc(file_path: str) -> str:
    """
    提取旧版 .doc 格式 Word 文档文本
    Windows: 优先使用 Word/WPS COM 自动化转换
    Linux/Docker: 优先使用 antiword，其次 LibreOffice
    最后兜底尝试 python-docx 兼容模式
    """
    import platform

    # 方式0（Windows）：使用 Word/WPS COM 自动化转换为 docx 再提取
    if platform.system() == "Windows":
        try:
            import win32com.client
            abs_path = os.path.abspath(file_path)
            with tempfile.TemporaryDirectory() as tmp_dir:
                docx_path = os.path.join(tmp_dir, "converted.docx")
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                try:
                    doc = word.Documents.Open(abs_path)
                    # SaveAs2 格式 16 = wdFormatDocumentDefault (.docx)
                    doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)
                    doc.Close(False)
                finally:
                    word.Quit()
                if os.path.exists(docx_path):
                    logger.info("[doc提取] 使用 Word/WPS COM 转换成功")
                    return _extract_docx(docx_path)
        except ImportError:
            logger.warning("[doc提取] pywin32 未安装，跳过 COM 方式（pip install pywin32）")
        except Exception as e:
            logger.warning(f"[doc提取] Word/WPS COM 转换失败: {e}")

    # 方式1：使用 antiword 提取文本（轻量级，Docker 环境推荐）
    if shutil.which("antiword"):
        try:
            result = subprocess.run(
                ["antiword", "-m", "UTF-8", file_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                logger.info(f"[doc提取] 使用 antiword 成功提取文本")
                return result.stdout.strip()
            else:
                logger.warning(f"[doc提取] antiword 返回码={result.returncode} stderr={result.stderr[:200]}")
        except subprocess.TimeoutExpired:
            logger.warning("[doc提取] antiword 执行超时")
        except Exception as e:
            logger.warning(f"[doc提取] antiword 异常: {e}")

    # 方式2：使用 LibreOffice 转换为 docx 再提取
    if shutil.which("libreoffice") or shutil.which("soffice"):
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                lo_cmd = "libreoffice" if shutil.which("libreoffice") else "soffice"
                subprocess.run(
                    [lo_cmd, "--headless", "--convert-to", "docx", "--outdir", tmp_dir, file_path],
                    capture_output=True, timeout=60,
                )
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                converted_path = os.path.join(tmp_dir, f"{base_name}.docx")
                if os.path.exists(converted_path):
                    logger.info(f"[doc提取] 使用 LibreOffice 转换成功")
                    return _extract_docx(converted_path)
        except Exception as e:
            logger.warning(f"[doc提取] LibreOffice 转换异常: {e}")

    # 方式3：尝试用 python-docx 直接打开（部分 .doc 文件实际是 XML 格式）
    try:
        logger.info(f"[doc提取] 尝试使用 python-docx 兼容模式")
        return _extract_docx(file_path)
    except Exception as e:
        logger.warning(f"[doc提取] python-docx 兼容模式失败: {e}")

    raise ValueError(
        "无法提取 .doc 文件内容。该文件可能是旧版 Word 二进制格式，"
        "建议使用 Word 或 WPS 将文件另存为 .docx 格式后重新上传。"
    )


def _extract_docx(file_path: str) -> str:
    """提取 Word 文档文本"""
    doc = DocxDocument(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


def _extract_pdf(file_path: str) -> str:
    """提取 PDF 文件文本"""
    doc = fitz.open(file_path)
    if doc.page_count > 100:
        doc.close()
        raise ValueError(f"PDF 页数超过限制（{doc.page_count}页，最多100页）")

    text_parts = []
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        if text.strip():
            text_parts.append(text.strip())

    doc.close()
    return "\n".join(text_parts)


def extract_html_from_file(file_path: str, file_ext: str, plain_text: str = "") -> str:
    """
    从文件中提取格式化 HTML，保留排版和字体样式
    Word 文档完整保留样式，TXT/PDF 使用纯文本包装

    :param file_path: 文件路径
    :param file_ext: 文件扩展名（含点，如 .docx）
    :param plain_text: 已提取的纯文本（用于 txt/pdf 的 HTML 包装）
    :return: HTML 字符串
    """
    ext = file_ext.lower()
    if ext == ".docx":
        return _extract_docx_html(file_path)
    elif ext == ".doc":
        # .doc 格式无法直接提取富文本样式，降级为纯文本 HTML 包装
        escaped = html_escape(plain_text)
        return f'<div style="white-space:pre-wrap;line-height:1.8;font-size:14px;">{escaped}</div>'
    else:
        escaped = html_escape(plain_text)
        return f'<div style="white-space:pre-wrap;line-height:1.8;font-size:14px;">{escaped}</div>'


def _extract_docx_html(file_path: str) -> str:
    """
    将 Word 文档转换为 HTML，保留排版和字体样式
    支持：段落样式、标题层级、对齐方式、首行缩进、段间距、行间距、
          加粗、斜体、下划线、删除线、字号、字体、颜色、表格
    """
    doc = DocxDocument(file_path)
    html_parts = []

    # 按文档顺序处理段落和表格
    para_idx = 0
    table_idx = 0

    for child in doc.element.body:
        tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag_name == 'p':
            if para_idx < len(doc.paragraphs):
                html_parts.append(_para_to_html(doc.paragraphs[para_idx]))
                para_idx += 1
        elif tag_name == 'tbl':
            if table_idx < len(doc.tables):
                html_parts.append(_table_to_html(doc.tables[table_idx]))
                table_idx += 1

    return '\n'.join(html_parts)


def _para_to_html(para) -> str:
    """将段落转换为 HTML 标签，保留排版样式"""
    # 空段落保留为空行
    if not para.text.strip():
        return '<p style="margin:0.3em 0;min-height:1em;"><br/></p>'

    # 根据样式确定标签（标题 → h1-h6，正文 → p）
    tag = 'p'
    try:
        style_name = (para.style.name or '') if para.style else ''
        for i in range(1, 7):
            if f'Heading {i}' in style_name or style_name == f'Heading{i}':
                tag = f'h{i}'
                break
    except Exception:
        pass

    # 收集段落样式
    styles = ['margin:0.3em 0']

    # 对齐方式
    try:
        if para.alignment is not None:
            align_val = int(para.alignment)
            align_map = {1: 'center', 2: 'right', 3: 'justify', 4: 'justify', 5: 'justify'}
            align = align_map.get(align_val, '')
            if align:
                styles.append(f'text-align:{align}')
    except Exception:
        pass

    # 首行缩进
    try:
        pf = para.paragraph_format
        if pf.first_line_indent and pf.first_line_indent.pt > 0:
            styles.append(f'text-indent:{pf.first_line_indent.pt:.1f}pt')
    except Exception:
        pass

    # 左缩进
    try:
        pf = para.paragraph_format
        if pf.left_indent and pf.left_indent.pt > 0:
            styles.append(f'padding-left:{pf.left_indent.pt:.1f}pt')
    except Exception:
        pass

    # 段前段后间距
    try:
        pf = para.paragraph_format
        if pf.space_before and pf.space_before.pt:
            styles.append(f'margin-top:{pf.space_before.pt:.1f}pt')
        if pf.space_after and pf.space_after.pt:
            styles.append(f'margin-bottom:{pf.space_after.pt:.1f}pt')
    except Exception:
        pass

    # 行间距
    try:
        pf = para.paragraph_format
        if pf.line_spacing is not None:
            if isinstance(pf.line_spacing, (int, float)):
                styles.append(f'line-height:{pf.line_spacing}')
            elif hasattr(pf.line_spacing, 'pt') and pf.line_spacing.pt:
                styles.append(f'line-height:{pf.line_spacing.pt:.1f}pt')
    except Exception:
        pass

    # 构建行内 HTML（保留字体样式）
    inline_html = _runs_to_html(para.runs)
    if not inline_html:
        inline_html = html_escape(para.text)

    style_attr = f' style="{";".join(styles)}"' if styles else ''
    return f'<{tag}{style_attr}>{inline_html}</{tag}>'


def _runs_to_html(runs) -> str:
    """将文档 Runs 转换为行内 HTML，保留字体样式"""
    parts = []
    for run in runs:
        text = run.text
        if not text:
            continue

        # HTML 转义
        text = html_escape(text)

        # 收集行内样式
        run_styles = []

        # 字体名称
        try:
            if run.font.name:
                run_styles.append(f'font-family:"{run.font.name}"')
        except Exception:
            pass

        # 字号
        try:
            if run.font.size and run.font.size.pt:
                run_styles.append(f'font-size:{run.font.size.pt:.1f}pt')
        except Exception:
            pass

        # 字体颜色
        try:
            if run.font.color and run.font.color.rgb:
                run_styles.append(f'color:#{run.font.color.rgb}')
        except Exception:
            pass

        # 加粗
        try:
            if run.bold:
                text = f'<strong>{text}</strong>'
        except Exception:
            pass

        # 斜体
        try:
            if run.italic:
                text = f'<em>{text}</em>'
        except Exception:
            pass

        # 下划线
        try:
            if run.underline:
                text = f'<u>{text}</u>'
        except Exception:
            pass

        # 删除线
        try:
            if run.font.strike:
                text = f'<s>{text}</s>'
        except Exception:
            pass

        # 包裹行内样式
        if run_styles:
            style_str = ';'.join(run_styles)
            text = f'<span style="{style_str}">{text}</span>'

        parts.append(text)

    return ''.join(parts)


def _table_to_html(table) -> str:
    """将表格转换为 HTML"""
    html = '<table style="border-collapse:collapse;width:100%;margin:8px 0;">'
    for row in table.rows:
        html += '<tr>'
        for cell in row.cells:
            cell_text = html_escape(cell.text.strip())
            html += f'<td style="border:1px solid #ccc;padding:6px 8px;vertical-align:top;">{cell_text}</td>'
        html += '</tr>'
    html += '</table>'
    return html


def generate_corrected_docx(
    original_path: str,
    issues: List[dict],
    output_path: str,
) -> str:
    """
    基于校对结果生成修订版 Word 文档
    在原文基础上用红色标记修改内容

    :param original_path: 原始 docx 文件路径
    :param issues: 校对问题列表
    :param output_path: 输出文件路径
    :return: 输出文件路径
    """
    doc = DocxDocument(original_path)

    # 构建替换映射
    replacements = {}
    for issue in issues:
        original = issue.get("original", "")
        suggestion = issue.get("suggestion", "")
        if original and suggestion and original != suggestion:
            replacements[original] = suggestion

    # 遍历段落进行替换
    for para in doc.paragraphs:
        full_text = para.text
        modified = False
        for original, suggestion in replacements.items():
            if original in full_text:
                full_text = full_text.replace(original, suggestion)
                modified = True

        if modified:
            # 清除原有 run，用新文本重建
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = full_text
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                run = para.add_run(full_text)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.save(output_path)
    logger.info(f"修订文档已生成: {output_path}")
    return output_path


def generate_corrected_txt(
    original_text: str,
    issues: List[dict],
    output_path: str,
) -> str:
    """
    基于校对结果生成修订版 TXT 文件

    :param original_text: 原始文本
    :param issues: 校对问题列表
    :param output_path: 输出文件路径
    :return: 输出文件路径
    """
    corrected = original_text
    for issue in issues:
        original = issue.get("original", "")
        suggestion = issue.get("suggestion", "")
        if original and suggestion and original != suggestion:
            corrected = corrected.replace(original, suggestion)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(corrected)

    logger.info(f"修订文本已生成: {output_path}")
    return output_path
